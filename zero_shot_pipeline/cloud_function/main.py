import functions_framework
import os
import uuid
import json
import re
# Removed openslide to avoid C dependencies on standard GCF
import numpy as np
import tifffile
import io
import base64
from PIL import Image
from docx import Document

import torch
import timm
from torchvision import transforms
from huggingface_hub import login

from google.cloud import aiplatform, storage, bigquery, secretmanager

# Configuration from Environment Variables
PROJECT_ID = os.environ.get("PROJECT_ID", "gsk-cmc-hackathon")
REGION = os.environ.get("REGION", "us-central1")
BUCKET_NAME = os.environ.get("BUCKET_NAME", f"{PROJECT_ID}-wsi-data")

# Model Endpoints & Secrets
MEDGEMMA_ID = os.environ.get("MEDGEMMA_ID", "YOUR_MEDGEMMA_ID")
MEDGEMMA_REGION = os.environ.get("MEDGEMMA_REGION", "us-east1")
MEDSIGLIP_ID = os.environ.get("MEDSIGLIP_ID", "YOUR_MEDSIGLIP_ID")
PATHFOUNDATION_ID = os.environ.get("PATHFOUNDATION_ID", "YOUR_PATHFOUNDATION_ID")
HF_TOKEN_SECRET = os.environ.get("HF_TOKEN_SECRET", "huggingface-token")

# Vector Search & Database
VECTOR_INDEX_ENDPOINT_ID = os.environ.get("VECTOR_INDEX_ENDPOINT_ID", "YOUR_VECTOR_ENDPOINT_ID")
VECTOR_INDEX_ID = os.environ.get("VECTOR_INDEX_ID", "YOUR_VECTOR_INDEX_ID")
BQ_TABLE = os.environ.get("BQ_TABLE", f"{PROJECT_ID}.pathology_db.tile_metadata")


@functions_framework.cloud_event
def trigger_pipeline(cloud_event):
    data = cloud_event.data
    file_name = data["name"]
    bucket_name = data["bucket"]
    
    # We only care about the inputs directory
    if not file_name.startswith('inputs/'):
        return
    print(f"Processing file {file_name} in bucket {bucket_name}")
    storage_client = storage.Client(project=PROJECT_ID)
    bucket = storage_client.bucket(bucket_name)
    uploaded_blob = bucket.blob(file_name)
    
    # ---------------------------------------------------------
    # 1. DUPLICATE HANDLING: Check if already processed
    # ---------------------------------------------------------
    uploaded_blob.reload()
    if uploaded_blob.metadata and uploaded_blob.metadata.get("pipeline_triggered") == "true":
        print(f"🛑 Duplicate trigger detected for {file_name}. This file has already been processed. Skipping.")
        return

    # ---------------------------------------------------------
    # 2. PAIRED FILE CHECK: Ensure both WSI and JSON exist
    # ---------------------------------------------------------
    if file_name.endswith('.svs'):
        wsi_blob = uploaded_blob
        wsi_name = file_name
        meta_name = file_name.replace('.svs', '_metadata.json')
        meta_blob = bucket.blob(meta_name)
    elif file_name.endswith('_metadata.json'):
        meta_blob = uploaded_blob
        meta_name = file_name
        wsi_name = file_name.replace('_metadata.json', '.svs')
        wsi_blob = bucket.blob(wsi_name)
    else:
        return # Ignore random files

    if not wsi_blob.exists():
        print(f"⏳ Waiting for pair: {wsi_name} is missing. Processing will not fire until WSI is uploaded.")
        return
        
    if not meta_blob.exists():
        print(f"⏳ Waiting for pair: {meta_name} is missing. Processing will not fire until JSON metadata is uploaded.")
        return

    # ---------------------------------------------------------
    # 3. MARK AS PROCESSED & PREPARE RUN
    # ---------------------------------------------------------
    wsi_blob.metadata = {"pipeline_triggered": "true"}
    wsi_blob.patch()
    meta_blob.metadata = {"pipeline_triggered": "true"}
    meta_blob.patch()

    print(f"✅ Both files detected ({wsi_blob.name} & {meta_blob.name}). Starting local pipeline execution!")
    
    aiplatform.init(project=PROJECT_ID, location=REGION)
    run_id = f"run_{uuid.uuid4().hex[:6]}"
    
    # ---------------------------------------------------------
    # STEP 1: Extract Metadata and Targets (MedGemma)
    # ---------------------------------------------------------
    print("🔍 [Step 1/3] Extracting metadata and prompting MedGemma...")
    metadata = run_extract_metadata_and_targets(meta_blob)
    print(f"📌 Targets to segment: {metadata.get('identified_targets', [])}")

    # ---------------------------------------------------------
    # STEP 2: Process WSI, Extract Embeddings, Segment & Index
    # ---------------------------------------------------------
    print("🔬 [Step 2/3] Downloading WSI, generating embeddings, segmenting & indexing...")
    wsi_gcs_uri = f"gs://{bucket_name}/{wsi_blob.name}"
    tile_count = run_process_wsi_and_index(wsi_gcs_uri, metadata, run_id, bucket_name)
    print(f"✅ Tiling & Indexing complete. Total tiles processed: {tile_count}")

    # ---------------------------------------------------------
    # STEP 3: Generate Report
    # ---------------------------------------------------------
    print("📄 [Step 3/3] Generating pathology summary report...")
    run_generate_report(run_id, metadata, tile_count, bucket_name)
    print("🎉 Success! Pipeline execution finished successfully.")


def run_extract_metadata_and_targets(meta_blob) -> dict:
    storage_client = storage.Client(project=PROJECT_ID)
    
    # 1. Safely Download JSON
    try:
        metadata_text = meta_blob.download_as_text()
        raw_json = json.loads(metadata_text)
    except Exception as e:
        raise RuntimeError(f"❌ CRITICAL ERROR: Failed to download or parse JSON metadata: {e}")
    
    # 2. FLEXIBLE JSON NORMALIZER
    flat_metadata = {}
    for item in raw_json.get("fieldValues", []):
        raw_key = item.get("name", "unknown_key")
        normalized_key = re.sub(r'[\s\-]+', '_', str(raw_key)).strip().upper()
        flat_metadata[normalized_key] = item.get("value", "")
        
    # Keep custom model configuration if provided in metadata JSON
    hf_model_id = raw_json.get("hf_model_id", os.environ.get("HF_MODEL_ID", "bioptimus/H-optimus-0"))
    flat_metadata["hf_model_id"] = hf_model_id
    
    # 3. Extract standard keys
    species = flat_metadata.get("SPECIES", "Unknown species")
    tissue = flat_metadata.get("TISSUE_NAME", "Unknown tissue")
    diagnosis = flat_metadata.get("PATIENT_PRIMARY_DIAGNOSIS", "Unknown diagnosis")
    disease_cat = flat_metadata.get("TISSUE_DISEASE_CATEGORY", "")
    
    prompt = f"As an expert pathologist, analyze a {species} {tissue} tissue sample. The patient's primary diagnosis is {diagnosis} ({disease_cat}). List 2-3 specific microscopic histological features, structures, or cellular abnormalities that should be segmented in this Whole Slide Image. Return ONLY a comma-separated list."
    
    # 4. Safely Call MedGemma
    try:
        endpoint = aiplatform.Endpoint(endpoint_name=MEDGEMMA_ID, location=MEDGEMMA_REGION)
        response = endpoint.predict(instances=[{"prompt": prompt}])
        targets = [t.strip() for t in response.predictions[0].split(',')]
    except Exception as e:
        print(f"⚠️ MedGemma Warning: Call failed ({e}). Falling back to default targets.")
        targets = ["Tumour Epithelium", "Necrosis", "Normal Adjacent Tissue"]
    
    flat_metadata["identified_targets"] = targets
    return flat_metadata


def run_process_wsi_and_index(wsi_gcs_uri: str, metadata: dict, run_id: str, bucket_name: str) -> int:
    print(f"🚩 [CHECKPOINT] Starting run_process_wsi_and_index for WSI: {wsi_gcs_uri}")
    storage_client = storage.Client(project=PROJECT_ID)
    bq_client = bigquery.Client(project=PROJECT_ID)
    
    # Initialize Vertex Endpoints
    print("🚩 [CHECKPOINT] Connecting to MedSigLIP and MatchingEngine endpoints...")
    ms_endpoint = aiplatform.Endpoint(MEDSIGLIP_ID)
    v_endpoint = aiplatform.MatchingEngineIndexEndpoint(VECTOR_INDEX_ENDPOINT_ID)
    
    # Check if visual embedding model is running via Vertex AI Model Garden endpoint
    use_vertex_pf = bool(PATHFOUNDATION_ID and PATHFOUNDATION_ID != "" and not PATHFOUNDATION_ID.startswith("YOUR_"))
    
    if use_vertex_pf:
        print(f"🚀 Running Path Foundation visual embeddings via Vertex AI Model Garden endpoint: {PATHFOUNDATION_ID}")
        pf_endpoint = aiplatform.Endpoint(PATHFOUNDATION_ID)
    else:
        print("🧠 PATHFOUNDATION_ID not set or invalid. Falling back to local HuggingFace PyTorch model on CPU...")
        hf_model_id = metadata.get("hf_model_id", os.environ.get("HF_MODEL_ID", "bioptimus/H-optimus-0"))
        
        # Fetch HF Token from Secret Manager
        try:
            print(f"🚩 [CHECKPOINT] Fetching HF Secret Token: {HF_TOKEN_SECRET}...")
            sm_client = secretmanager.SecretManagerServiceClient()
            secret_name = f"projects/{PROJECT_ID}/secrets/{HF_TOKEN_SECRET}/versions/latest"
            hf_token = sm_client.access_secret_version(request={"name": secret_name}).payload.data.decode("UTF-8")
            login(token=hf_token)
        except Exception as e:
            print(f"⚠️ Warning: Failed to fetch HF token {HF_TOKEN_SECRET}: {e}. Proceeding without auth (will fail for gated models).")
        
        # Load Model (CPU optimized since Cloud Functions don't run with GPU)
        device = "cpu"
        print(f"🚩 [CHECKPOINT] Loading Timm model '{hf_model_id}' on {device}...")
        
        if "Virchow" in hf_model_id:
            from timm.layers import SwiGLUPacked
            model = timm.create_model(f"hf-hub:{hf_model_id}", pretrained=True, mlp_layer=SwiGLUPacked, act_layer=torch.nn.SiLU, dynamic_img_size=False)
            from timm.data import resolve_data_config
            from timm.data.transforms_factory import create_transform
            transform = create_transform(**resolve_data_config(model.pretrained_cfg, model=model))
        else:
            # Default / H-optimus-0
            model = timm.create_model(f"hf-hub:{hf_model_id}", pretrained=True, init_values=1e-5, dynamic_img_size=False)
            transform = transforms.Compose([
                transforms.ToTensor(),
                transforms.Normalize(mean=(0.707223, 0.578729, 0.703617), std=(0.211883, 0.230117, 0.177517)),
            ])
            
        model = model.to(device)
        model.eval()
        print("🚩 [CHECKPOINT] Local model loaded and set to eval successfully.")
    
    wsi_bucket = wsi_gcs_uri.split("/")[2]
    wsi_blob_name = "/".join(wsi_gcs_uri.split("/")[3:])
    wsi_filename = wsi_gcs_uri.split("/")[-1]
    
    processed_count = 0
    targets = metadata.get("identified_targets", ["Abnormal Tissue"])
    
    # Open slide directly from GCS stream in pure Python using tifffile (no /tmp download needed!)
    try:
        print(f"🚩 [CHECKPOINT] Opening binary stream for gs://{wsi_bucket}/{wsi_blob_name} directly from GCS...")
        bucket = storage_client.bucket(wsi_bucket)
        blob = bucket.blob(wsi_blob_name)
        
        # Open blob as seekable binary network stream
        gcs_file = blob.open("rb")
        print("🚩 [CHECKPOINT] Stream initialized. Parsing TIFF structures...")
        tif = tifffile.TiffFile(gcs_file)
        page = tif.pages[0]
        height, width = page.shape[:2]
        print(f"🚩 [CHECKPOINT] slide successfully parsed. Shape dimensions: {width}x{height}")
    except Exception as e:
        raise RuntimeError(f"❌ CRITICAL ERROR: Failed to stream and parse WSI file from GCS: {e}")
        
    try:
        print("🚩 [CHECKPOINT] Slicing slide into tiles and starting extraction loop...")
        # Extract 224x224 tiles to match native ViT resolution
        for x in range(0, width, 224):
            for y in range(0, height, 224):
                try:
                    # Extract the pixel window natively from the tiled TIFF page
                    tile_np = page.asarray(window=((y, y+224), (x, x+224)))
                    tile = Image.fromarray(tile_np).convert("RGB")
                except Exception as e:
                    print(f"⚠️ Warning: Failed to read tile at {x},{y}: {e}")
                    continue # Skip corrupted coordinate
            
            if tile.getextrema()[0][0] < 240: # Skip blank glass
                tile_id = f"tile_x{x}_y{y}"
                vector_id = f"vec_{run_id}_{tile_id}"
                print(f"🚩 [CHECKPOINT TILE] Sliced tile '{tile_id}' at x={x}, y={y}. Processing...")
                
                buffered = io.BytesIO()
                tile.save(buffered, format="PNG")
                img_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
                
                # Embedding extraction
                if use_vertex_pf:
                    try:
                        print(f"🚩 [CHECKPOINT TILE] Calling Model Garden Endpoint for tile {tile_id} embeddings...")
                        response = pf_endpoint.predict(instances=[{"image_bytes": img_b64}])
                        embedding = response.predictions[0]
                    except Exception as e:
                        print(f"⚠️ Warning: Model Garden Endpoint embedding failed on {tile_id}. Error: {e}")
                        continue
                else:
                    # Embedding extraction via local HF Model
                    try:
                        print(f"🚩 [CHECKPOINT TILE] Extracting local PyTorch embeddings for tile {tile_id}...")
                        with torch.inference_mode():
                            input_tensor = transform(tile).unsqueeze(0).to(device)
                            output = model(input_tensor)
                            
                            if "Virchow" in hf_model_id:
                                class_token = output[:, 0]
                                patch_tokens = output[:, 1:]
                                embedding_tensor = torch.cat([class_token, patch_tokens.mean(1)], dim=-1)
                            else:
                                embedding_tensor = output # H-optimus-0
                                
                            embedding = embedding_tensor.cpu().to(torch.float32).numpy().flatten().tolist()
                    except Exception as e:
                        print(f"⚠️ Warning: Local embedding generation failed on {tile_id}. Error: {e}")
                        continue
                
                # Vision API Segmentation Mask
                try:
                    print(f"🚩 [CHECKPOINT TILE] Calling MedSigLIP endpoint to generate mask for tile {tile_id}...")
                    mask = np.array(ms_endpoint.predict(instances=[{"image_bytes": img_b64, "targets": targets}]).predictions[0], dtype=np.uint8) 
                except Exception as e:
                    print(f"⚠️ Warning: MedSigLIP segmentation failed on {tile_id}. Error: {e}. Skipping.")
                    continue
                
                # Storage & Database Error Handling
                try:
                    print(f"🚩 [CHECKPOINT TILE] Saving mask, upserting vector index, and logging to BigQuery for {tile_id}...")
                    local_mask = f"/tmp/{tile_id}_mask.tif"
                    tifffile.imwrite(local_mask, mask, photometric='minisblack')
                    mask_uri = f"gs://{wsi_bucket}/outputs/{run_id}/masks/{tile_id}_mask.tif"
                    storage_client.bucket(wsi_bucket).blob(f"outputs/{run_id}/masks/{tile_id}_mask.tif").upload_from_filename(local_mask)
                    
                    v_endpoint.upsert_datapoints(index_id=VECTOR_INDEX_ID, datapoints=[{"datapoint_id": vector_id, "feature_vector": embedding}])
                    
                    # Log normalized metadata to BigQuery
                    rows = [{
                        "run_id": run_id, "wsi_filename": wsi_filename,
                        "heart_id": metadata.get("HEART_ID", ""),
                        "tissue_name": metadata.get("TISSUE_NAME", ""),
                        "primary_diagnosis": metadata.get("PATIENT_PRIMARY_DIAGNOSIS", ""),
                        "sample_category": metadata.get("SAMPLE_CATEGORY", ""),
                        "stain": metadata.get("STAIN_FULL_NAME", ""),
                        "species": metadata.get("SPECIES", ""),
                        "disease_category": metadata.get("TISSUE_DISEASE_CATEGORY", ""),
                        "tile_id": tile_id, "vector_id": vector_id, 
                        "mask_uri": mask_uri, "predicted_class": targets[0]
                    }]
                    bq_client.insert_rows_json(BQ_TABLE, rows)
                    
                    processed_count += 1
                    if processed_count >= 20: 
                        print("🚩 [CHECKPOINT] Tile processed count capped at 20. Finalizing run...")
                        return processed_count # Cap for POC speed
                except Exception as e:
                    print(f"⚠️ Warning: Database/Storage operation failed on {tile_id}: {e}")
                    continue
                    
    finally:
        print("🚩 [CHECKPOINT] Releasing file resources and closing network streams...")
        tif.close()
        gcs_file.close()
        
    return processed_count


def run_generate_report(run_id: str, metadata: dict, tile_count: int, bucket_name: str):
    doc = Document()
    doc.add_heading('WSI Zero-Shot Processing Summary', 0)
    doc.add_paragraph(f"Run ID: {run_id}")
    doc.add_paragraph(f"HEART ID: {metadata.get('HEART_ID', 'N/A')}")
    doc.add_paragraph(f"Diagnosis: {metadata.get('PATIENT_PRIMARY_DIAGNOSIS', 'N/A')}")
    doc.add_paragraph(f"Tissue / Species: {metadata.get('TISSUE_NAME', 'N/A')} ({metadata.get('SPECIES', 'N/A')})")
    doc.add_paragraph(f"Stain: {metadata.get('STAIN_FULL_NAME', 'N/A')}")
    doc.add_paragraph(f"AI Targets Segmented: {', '.join(metadata.get('identified_targets', []))}")
    doc.add_paragraph(f"Total Tiles Processed & Indexed: {tile_count}")
    
    local_doc = f"/tmp/{run_id}_summary.docx"
    doc.save(local_doc)
    
    storage_client = storage.Client(project=PROJECT_ID)
    storage_client.bucket(bucket_name).blob(f"outputs/{run_id}/{run_id}_summary.docx").upload_from_filename(local_doc)
