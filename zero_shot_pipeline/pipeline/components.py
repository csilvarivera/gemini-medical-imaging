from kfp.v2.dsl import component
import os

# Base image should be provided at compilation or execution time. 
# We'll set a placeholder or use an environment variable for the decorator if possible,
# or pass it as a parameter if needed. In KFP v2, `base_image` must be a literal string 
# at compile time. We assume standard python:3.10 and let the components install what they need 
# or hardcode the expected container registry path.
BASE_IMAGE = "us-central1-docker.pkg.dev/gsk-cmc-hackathon/histology-repo/kfp-base:latest"

@component(base_image=BASE_IMAGE)
def extract_metadata_and_targets(metadata_gcs_uri: str, medgemma_id: str, project: str, region: str) -> dict:
    import json
    import re
    from google.cloud import storage, aiplatform
    
    aiplatform.init(project=project, location=region)
    storage_client = storage.Client(project=project)
    
    # 1. Safely Download JSON
    try:
        bucket_name = metadata_gcs_uri.split("/")[2]
        blob_path = "/".join(metadata_gcs_uri.split("/")[3:])
        metadata_text = storage_client.bucket(bucket_name).blob(blob_path).download_as_text()
        raw_json = json.loads(metadata_text)
    except Exception as e:
        raise RuntimeError(f"❌ CRITICAL ERROR: Failed to download or parse JSON metadata: {e}")
    
    # 2. FLEXIBLE JSON NORMALIZER
    # Converts "TissueName", "tissue_name", "Tissue Name" all into "TISSUE_NAME"
    flat_metadata = {}
    for item in raw_json.get("fieldValues", []):
        raw_key = item.get("name", "unknown_key")
        normalized_key = re.sub(r'[\s\-]+', '_', str(raw_key)).strip().upper()
        flat_metadata[normalized_key] = item.get("value", "")
    
    # 3. Extract standard keys
    species = flat_metadata.get("SPECIES", "Unknown species")
    tissue = flat_metadata.get("TISSUE_NAME", "Unknown tissue")
    diagnosis = flat_metadata.get("PATIENT_PRIMARY_DIAGNOSIS", "Unknown diagnosis")
    disease_cat = flat_metadata.get("TISSUE_DISEASE_CATEGORY", "")
    
    prompt = f"As an expert pathologist, analyze a {species} {tissue} tissue sample. The patient's primary diagnosis is {diagnosis} ({disease_cat}). List 2-3 specific microscopic histological features, structures, or cellular abnormalities that should be segmented in this Whole Slide Image. Return ONLY a comma-separated list."
    
    # 4. Safely Call MedGemma
    try:
        endpoint = aiplatform.Endpoint(medgemma_id)
        response = endpoint.predict(instances=[{"prompt": prompt}])
        targets = [t.strip() for t in response.predictions[0].split(',')]
    except Exception as e:
        print(f"⚠️ MedGemma Warning: Call failed ({e}). Falling back to default targets.")
        targets = ["Tumour Epithelium", "Necrosis", "Normal Adjacent Tissue"]
    
    flat_metadata["identified_targets"] = targets
    return flat_metadata


@component(base_image=BASE_IMAGE)
def process_wsi_and_index(
    wsi_gcs_uri: str, metadata: dict, run_id: str,
    pf_id: str, ms_id: str, v_end_id: str, v_idx_id: str, bq_table: str, project: str, region: str
) -> int:
    import openslide, numpy as np, tifffile, os, base64, io
    from google.cloud import storage, aiplatform, bigquery
    
    aiplatform.init(project=project, location=region)
    storage_client = storage.Client(project=project)
    bq_client = bigquery.Client(project=project)
    
    pf_endpoint = aiplatform.Endpoint(pf_id)
    ms_endpoint = aiplatform.Endpoint(ms_id)
    v_endpoint = aiplatform.MatchingEngineIndexEndpoint(v_end_id)
    
    wsi_bucket = wsi_gcs_uri.split("/")[2]
    wsi_filename = wsi_gcs_uri.split("/")[-1]
    local_wsi_path = f"/tmp/{wsi_filename}"
    
    # Safely download massive WSI
    try:
        storage_client.bucket(wsi_bucket).blob("/".join(wsi_gcs_uri.split("/")[3:])).download_to_filename(local_wsi_path)
        slide = openslide.OpenSlide(local_wsi_path)
    except Exception as e:
        raise RuntimeError(f"❌ CRITICAL ERROR: Failed to download or open WSI file: {e}")
        
    width, height = slide.dimensions
    processed_count = 0
    targets = metadata.get("identified_targets", ["Abnormal Tissue"])
    
    for x in range(0, width, 512):
        for y in range(0, height, 512):
            try:
                tile = slide.read_region((x, y), 0, (512, 512)).convert("RGB")
            except Exception as e:
                print(f"⚠️ Warning: Failed to read tile at {x},{y}: {e}")
                continue # Skip corrupted coordinate
            
            if tile.getextrema()[0][0] < 240: # Skip blank glass
                tile_id = f"tile_x{x}_y{y}"
                vector_id = f"vec_{run_id}_{tile_id}"
                
                buffered = io.BytesIO()
                tile.save(buffered, format="PNG")
                img_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
                
                # Vision API Error Handling
                try:
                    embedding = pf_endpoint.predict(instances=[{"image_bytes": img_b64}]).predictions[0] 
                    mask = np.array(ms_endpoint.predict(instances=[{"image_bytes": img_b64, "targets": targets}]).predictions[0], dtype=np.uint8) 
                except Exception as e:
                    print(f"⚠️ Warning: Vision AI failed on {tile_id}. Error: {e}. Skipping.")
                    continue
                
                # Storage & Database Error Handling
                try:
                    local_mask = f"/tmp/{tile_id}_mask.tif"
                    tifffile.imwrite(local_mask, mask, photometric='minisblack')
                    mask_uri = f"gs://{wsi_bucket}/outputs/{run_id}/masks/{tile_id}_mask.tif"
                    storage_client.bucket(wsi_bucket).blob(f"outputs/{run_id}/masks/{tile_id}_mask.tif").upload_from_filename(local_mask)
                    
                    v_endpoint.upsert_datapoints(index_id=v_idx_id, datapoints=[{"datapoint_id": vector_id, "feature_vector": embedding}])
                    
                    # Log normalized metadata to BigQuery
                    rows = [{
                        "run_id": run_id, "wsi_filename": wsi_filename,
                        "heart_id": metadata.get("HEART_ID", ""),
                        "tissue_name": metadata.get("TISSUE_NAME", ""),
                        "primary_diagnosis": metadata.get("PATIENT_PRIMARY_DIAGNOSIS", ""),
                        "sample_category": metadata.get("SAMPLE_CATEGORY", ""),
                        "stain": metadata.get("STAIN_FULL_NAME", ""),
                        "species": metadata.get("SPECIES", ""),
                        "disease_category": metadata.get("TISSUE_DISEASE_CATEGORY", ""),
                        "tile_id": tile_id, "vector_id": vector_id, 
                        "mask_uri": mask_uri, "predicted_class": targets[0]
                    }]
                    bq_client.insert_rows_json(bq_table, rows)
                    
                    processed_count += 1
                    if processed_count >= 20: return processed_count # Cap for POC speed
                except Exception as e:
                    print(f"⚠️ Warning: Database/Storage operation failed on {tile_id}: {e}")
                    continue
                    
    return processed_count


@component(base_image=BASE_IMAGE)
def generate_report(run_id: str, metadata: dict, tile_count: int, output_bucket: str, project: str):
    from docx import Document
    from google.cloud import storage
    
    doc = Document()
    doc.add_heading('WSI Zero-Shot Processing Summary', 0)
    doc.add_paragraph(f"Run ID: {run_id}")
    doc.add_paragraph(f"HEART ID: {metadata.get('HEART_ID', 'N/A')}")
    doc.add_paragraph(f"Diagnosis: {metadata.get('PATIENT_PRIMARY_DIAGNOSIS', 'N/A')}")
    doc.add_paragraph(f"Tissue / Species: {metadata.get('TISSUE_NAME', 'N/A')} ({metadata.get('SPECIES', 'N/A')})")
    doc.add_paragraph(f"Stain: {metadata.get('STAIN_FULL_NAME', 'N/A')}")
    doc.add_paragraph(f"AI Targets Segmented: {', '.join(metadata.get('identified_targets', []))}")
    doc.add_paragraph(f"Total Tiles Processed & Indexed: {tile_count}")
    
    local_doc = f"/tmp/{run_id}_summary.docx"
    doc.save(local_doc)
    storage.Client(project=project).bucket(output_bucket).blob(f"outputs/{run_id}/{run_id}_summary.docx").upload_from_filename(local_doc)
